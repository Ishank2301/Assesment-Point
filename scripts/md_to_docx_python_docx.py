"""Convert markdown report to DOCX using python-docx for richer formatting."""

import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
import json


def add_paragraph_with_inline_formatting(doc, text: str):
    # Supports **bold** and *italic* minimally
    parts = re.split(r"(\*\*.*?\*\*|\*.*?\*)", text)
    p = doc.add_paragraph()
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


def convert(md_path: Path, docx_path: Path):
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    doc = Document()
    # Basic styling
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    if "Heading 1" in doc.styles:
        h1 = doc.styles["Heading 1"]
        h1.font.name = "Calibri"
        h1.font.size = Pt(16)
        h1.font.bold = True
    in_code = False
    code_lines = []
    bullet_mode = False

    for raw in lines:
        line = raw.rstrip("\n")
        if line.strip().startswith("```"):
            if in_code:
                # flush code block
                p = doc.add_paragraph()
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                in_code = False
                code_lines = []
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue

        if line.strip().startswith("#"):
            # heading level
            m = re.match(r"^(#+)\s+(.*)$", line.strip())
            if m:
                level = min(len(m.group(1)), 6)
                doc.add_heading(m.group(2), level=level)
            else:
                doc.add_paragraph(line)
            bullet_mode = False
            continue

        if line.strip().startswith("- "):
            doc.add_paragraph(line.strip()[2:], style="List Bullet")
            bullet_mode = True
            continue

        if not line.strip():
            bullet_mode = False
            doc.add_paragraph("")
            continue

        # normal paragraph with inline formatting
        add_paragraph_with_inline_formatting(doc, line)

    doc.save(docx_path)
    # Append metrics and figures if available
    root = md_path.resolve().parents[1]
    metrics_path = root / "reports" / "metrics.json"
    fi_path = root / "reports" / "figures" / "feature_importance.png"
    dec_chart = root / "scorer_results" / "candidate_december.png"
    # Reopen to append
    doc = Document(docx_path)
    doc.add_page_break()
    doc.add_heading("Results and Figures", level=1)
    if metrics_path.exists():
        with metrics_path.open(encoding="utf-8") as f:
            metrics = json.load(f)
        table = doc.add_table(rows=1, cols=4)
        hdr = table.rows[0].cells
        hdr[0].text = "Metric"
        hdr[1].text = "Train"
        hdr[2].text = "Validation"
        hdr[3].text = "Unit"

        def add_row(name, train, val, unit=""):
            r = table.add_row().cells
            r[0].text = name
            r[1].text = f"{train:.2f}"
            r[2].text = f"{val:.2f}"
            r[3].text = unit

        add_row("MAE", metrics["train"]["mae"], metrics["val"]["mae"], "")
        add_row("RMSE", metrics["train"]["rmse"], metrics["val"]["rmse"], "")
        add_row("MAPE", metrics["train"]["mape"], metrics["val"]["mape"], "%")
    if fi_path.exists():
        doc.add_paragraph("Feature importance:")
        doc.add_picture(str(fi_path), width=Inches(6))
    if dec_chart.exists():
        doc.add_paragraph("December candidate chart:")
        doc.add_picture(str(dec_chart), width=Inches(6))
    doc.save(docx_path)


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[1]
    md = root / "reports" / "model_report.md"
    out = root / "reports" / "model_report.docx"
    if not md.exists():
        print("Markdown not found:", md)
    else:
        convert(md, out)
        print("Wrote", out)
